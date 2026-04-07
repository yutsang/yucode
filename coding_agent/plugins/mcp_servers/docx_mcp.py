"""MCP server for Word (.docx) file operations.

Requires: python-docx >= 1.1

Run:  python -m coding_agent.plugins.mcp_servers.docx_mcp
"""

from __future__ import annotations

from typing import Any

from ._protocol import McpStdioServer


def _require_docx():
    try:
        import docx  # noqa: F401
        return docx
    except ImportError as exc:
        raise RuntimeError("python-docx is required: pip install python-docx>=1.1") from exc


def handle_read_paragraphs(args: dict[str, Any]) -> Any:
    docx = _require_docx()
    doc = docx.Document(args["path"])
    max_paragraphs = int(args.get("max_paragraphs", 1000))
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        if i >= max_paragraphs:
            break
        paragraphs.append({"index": i, "style": para.style.name, "text": para.text})
    return {"paragraphs": paragraphs, "total": len(doc.paragraphs)}


def handle_read_tables(args: dict[str, Any]) -> Any:
    docx = _require_docx()
    doc = docx.Document(args["path"])
    tables = []
    for t_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append({"index": t_idx, "rows": rows})
    return {"tables": tables}


def handle_read_full_text(args: dict[str, Any]) -> Any:
    docx = _require_docx()
    doc = docx.Document(args["path"])
    text = "\n".join(p.text for p in doc.paragraphs)
    max_chars = int(args.get("max_chars", 50000))
    if len(text) > max_chars:
        text = text[:max_chars] + f"\n... truncated at {max_chars} chars"
    return text


def handle_append_text(args: dict[str, Any]) -> Any:
    docx = _require_docx()
    path = args["path"]
    doc = docx.Document(path)
    style = args.get("style", "Normal")
    doc.add_paragraph(args["text"], style=style)
    doc.save(path)
    return {"appended": True, "style": style}


def handle_write_new(args: dict[str, Any]) -> Any:
    docx = _require_docx()
    path = args["path"]
    doc = docx.Document()
    for para in args.get("paragraphs", []):
        text = para if isinstance(para, str) else str(para.get("text", ""))
        style = "Normal" if isinstance(para, str) else str(para.get("style", "Normal"))
        doc.add_paragraph(text, style=style)
    doc.save(path)
    return {"created": path}


def main():
    server = McpStdioServer("yucode-docx")

    server.register_tool("read_paragraphs", "Read paragraphs from a Word document.", {
        "type": "object",
        "properties": {
            "path": {"type": "string", "description": "Path to .docx file"},
            "max_paragraphs": {"type": "integer", "description": "Max paragraphs to read (default 1000)"},
        },
        "required": ["path"],
    }, handle_read_paragraphs)

    server.register_tool("read_tables", "Read all tables from a Word document.", {
        "type": "object",
        "properties": {"path": {"type": "string"}},
        "required": ["path"],
    }, handle_read_tables)

    server.register_tool("read_full_text", "Read the entire document as plain text.", {
        "type": "object",
        "properties": {
            "path": {"type": "string"},
            "max_chars": {"type": "integer", "description": "Max characters (default 50000)"},
        },
        "required": ["path"],
    }, handle_read_full_text)

    server.register_tool("append_text", "Append a paragraph to an existing Word document.", {
        "type": "object",
        "properties": {
            "path": {"type": "string"},
            "text": {"type": "string"},
            "style": {"type": "string", "description": "Paragraph style (default Normal)"},
        },
        "required": ["path", "text"],
    }, handle_append_text)

    server.register_tool("write_new", "Create a new Word document with paragraphs.", {
        "type": "object",
        "properties": {
            "path": {"type": "string"},
            "paragraphs": {
                "type": "array",
                "items": {"oneOf": [
                    {"type": "string"},
                    {"type": "object", "properties": {"text": {"type": "string"}, "style": {"type": "string"}}},
                ]},
            },
        },
        "required": ["path", "paragraphs"],
    }, handle_write_new)

    server.serve()


if __name__ == "__main__":
    main()
